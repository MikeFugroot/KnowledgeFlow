# -*- coding: utf-8 -*-
"""
文件解析服务 — 复用桌面版 read_txt/read_pdf/read_docx/read_video 逻辑
去掉所有 PyQt 依赖，改为纯 Python + asyncio 兼容
"""

import os
import re
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from app.config import settings
from app.utils.text import clean_text


def read_txt(path: Path) -> Dict[str, Any]:
    """读取纯文本文件"""
    encodings = ["utf-8", "gbk", "gb18030", "utf-16"]
    last_err = None
    for enc in encodings:
        try:
            text = path.read_text(encoding=enc)
            return {
                "type": "text",
                "path": str(path),
                "success": True,
                "title": path.stem,
                "text": clean_text(text),
                "error": "",
            }
        except Exception as e:
            last_err = e
    return {"type": "text", "path": str(path), "success": False, "title": path.stem, "text": "", "error": str(last_err)}


def read_docx(path: Path) -> Dict[str, Any]:
    """读取 Word 文档"""
    result = {"type": "docx", "path": str(path), "success": False, "title": path.stem, "text": "", "error": ""}
    try:
        import docx
        doc = docx.Document(str(path))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        # 读取表格文字
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        result["success"] = True
        result["text"] = clean_text("\n".join(parts))
    except Exception as e:
        result["error"] = str(e)
    return result


def read_pdf(path: Path) -> Dict[str, Any]:
    """读取 PDF 文件"""
    result = {"type": "pdf", "path": str(path), "success": False, "title": path.stem, "text": "", "pages": [], "error": ""}
    try:
        import fitz
        doc = fitz.open(str(path))
        all_text = []
        for page_index in range(len(doc)):
            page = doc[page_index]
            page_text = clean_text(page.get_text("text"))
            result["pages"].append({
                "page_number": page_index + 1,
                "char_count": len(page_text),
                "text_preview": page_text[:300],
            })
            if page_text:
                all_text.append(f"\n\n===== 第 {page_index + 1} 页 =====\n\n{page_text}")
        doc.close()
        result["success"] = True
        result["text"] = clean_text("\n".join(all_text))
    except Exception as e:
        result["error"] = str(e)
    return result


def read_video_with_faster_whisper(
    path: Path,
    log_func: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    """使用 faster-whisper 转写视频/音频"""
    result = {"type": "video", "path": str(path), "success": False, "title": path.stem, "text": "", "segments": [], "error": ""}
    try:
        import torch
        from faster_whisper import WhisperModel

        device = "cuda" if torch.cuda.is_available() else "cpu"
        compute_type = (
            settings.FASTER_WHISPER_GPU_COMPUTE_TYPE
            if device == "cuda"
            else settings.FASTER_WHISPER_CPU_COMPUTE_TYPE
        )
        last_error = None

        model_candidates = []
        if settings.USE_LOCAL_FASTER_WHISPER_FIRST:
            local_dir = Path(settings.LOCAL_FASTER_WHISPER_MODEL_DIR)
            if local_dir.exists():
                model_candidates.append(str(local_dir.resolve()))
            elif log_func:
                log_func(f"本地 faster-whisper 模型目录不存在: {local_dir}")
        model_candidates.extend(settings.whisper_model_candidates_list)

        for model_name in model_candidates:
            try:
                if log_func:
                    log_func(f"加载 faster-whisper 模型：{model_name}，device={device}")
                model = WhisperModel(model_name, device=device, compute_type=compute_type)
                segments_iter, info = model.transcribe(
                    str(path),
                    language=settings.WHISPER_LANGUAGE,
                    beam_size=settings.WHISPER_BEAM_SIZE,
                    vad_filter=True,
                    vad_parameters={"min_silence_duration_ms": 500},
                )

                segments = []
                texts = []
                for seg in segments_iter:
                    seg_text = clean_text(seg.text)
                    if seg_text:
                        texts.append(seg_text)
                    segments.append({
                        "start": round(float(seg.start), 2),
                        "end": round(float(seg.end), 2),
                        "text": seg_text,
                    })

                result["success"] = True
                result["text"] = clean_text(" ".join(texts))
                result["segments"] = segments
                result["model"] = model_name
                result["backend"] = "faster_whisper"
                return result
            except Exception as e:
                last_error = e
                if log_func:
                    log_func(f"模型 {model_name} 失败，尝试降级。原因：{e}")

        raise last_error or RuntimeError("faster-whisper 转写失败")

    except Exception as e:
        result["error"] = str(e)
    return result


def read_video_with_openai_whisper(
    path: Path,
    log_func: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    """使用 openai-whisper 转写视频/音频（兜底方案）"""
    result = {"type": "video", "path": str(path), "success": False, "title": path.stem, "text": "", "segments": [], "error": ""}
    try:
        import torch
        import whisper

        model_name = "medium"
        if log_func:
            log_func(f"加载 openai-whisper 模型：{model_name}")
        model = whisper.load_model(model_name)
        transcribe_result = model.transcribe(
            str(path),
            language=settings.WHISPER_LANGUAGE,
            fp16=torch.cuda.is_available(),
            beam_size=settings.WHISPER_BEAM_SIZE,
            verbose=False,
        )
        result["success"] = True
        result["text"] = clean_text(transcribe_result.get("text", ""))
        result["segments"] = [
            {
                "start": round(seg.get("start", 0), 2),
                "end": round(seg.get("end", 0), 2),
                "text": clean_text(seg.get("text", "")),
            }
            for seg in transcribe_result.get("segments", [])
        ]
        result["model"] = model_name
        result["backend"] = "openai_whisper"
    except Exception as e:
        result["error"] = str(e)
    return result


def read_video(
    path: Path,
    log_func: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    """读取视频/音频文件，自动选择转写引擎"""
    if settings.WHISPER_BACKEND == "faster_whisper":
        result = read_video_with_faster_whisper(path, log_func=log_func)
        if result.get("success"):
            return result
        if log_func:
            log_func(f"faster-whisper 失败，尝试 openai-whisper 兜底：{result.get('error')}")
    return read_video_with_openai_whisper(path, log_func=log_func)


def read_any_file(
    path: Path,
    log_func: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    """
    自动识别文件类型并读取，返回统一格式的解析结果

    返回格式：
    {
        "type": "pdf/docx/text/video/...",
        "path": "原始路径",
        "success": True/False,
        "title": "文件标题",
        "text": "提取的文本",
        "error": "错误信息",
        "pages": [...],      # PDF 专有
        "segments": [...],    # 视频 专有
    }
    """
    ext = path.suffix.lower()
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext in [".txt", ".md"]:
        return read_txt(path)
    if ext in settings.video_exts_set:
        return read_video(path, log_func=log_func)
    # 对 PPT/Excel/图片类文件，本地不做复杂解析，但允许后续 Qwen-Doc-Turbo 直接读原文件
    if ext in settings.qwen_doc_supported_exts_set:
        return {"type": ext.lstrip("."), "path": str(path), "success": True, "title": path.stem, "text": "", "error": ""}
    return {"type": "unknown", "path": str(path), "success": False, "title": path.stem, "text": "", "error": f"不支持的文件类型：{ext}"}
